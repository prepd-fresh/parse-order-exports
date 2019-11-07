"""
From the exported order file,
create the spreadsheet needed for further processing.
Also create the stickers for the meal bags.
"""
import sys
import csv
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

def extractCustomerData(fileName):
    """
    Extract the customer data from the spreadsheet
    """
    customers = []
    addresses = []

    weeklyOrders = defaultdict(list)

    customerIDs = {}

    with open(fileName) as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            customer_name = row['First Name (Billing)'] + " " + row['Last Name (Billing)']
            customer_address = row['Address 1&2 (Billing)']
            customer_id = customer_name + "@" + customer_address
            customer_order = row['Item Name'] + "*" + row['Quantity']

            weeklyOrders[customer_id].append(customer_order)
            customerIDs[customer_name + '@' + customer_address] = customer_order

    return customerIDs, weeklyOrders

def generateLabels(customerIds):
    """
    Generate the custom label names for the stickers
    """
    names = []
    addresses = []
    totalMeals = []

    for customer in customerIds:
        names.append(customer.split('@')[0])
        addresses.append(customer.split('@')[1])
        totalMeals.append(customerIds[customer])

    return names, addresses, totalMeals

def createStickers(customers, addresses, quantities, extra_text="Please Refrigerate!"):
    """
    Create the stickers that will be placed on the meal bags
    """
    document = Document('templates/BlankAddressLabels.docx')

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.19)
        section.left_margin = Cm(0.98)
        section.right_margin = Cm(0.79)


    table = document.add_table(rows=0, cols=4)

    EXTRA_TEXT = extra_text

    customerNameStickers = []
    customerAddressStickers = []

    for i in range(len(customers)):
        try:
            total = int(quantities[i])
        except ValueError:
            total = 0

        while int(total) > 0:
            customerNameStickers.append(customers[i])
            customerAddressStickers.append(addresses[i])
            quantities[i] = int(quantities[i]) - 5

    chunks = createChunks(15, customerNameStickers)

    addressDict = {}
    addressDict[""] = "null"
    for customer, address in zip(customers, addresses):
        addressDict[customer] = address

    row_cells = table.add_row().cells
    j = 0
    for i in range(15):
        for j in range(4):
            try:
                row_cells[j].text = chunks[j][i] + "\n" + addressDict[chunks[j][i]].split(',')[0] + "\n" + EXTRA_TEXT
                row_cells[j].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[j].paragraphs[0].paragraph_format.line_spacing = 1
            except IndexError:
                break
        row_cells = table.add_row().cells

        if ((i + 1) % 15 == 0):
            j += 1

    document.save('AddressLabels.docx')

def createChunks(n, names):
    """
    Create segmented chunks so that the names on the sticker labels
    will line up how we want them to
    """
    chunks = []
    currentChunk = []

    for i in range(len(names)):
        if ((i) % n == 0 and i != 0):
            chunks.append(currentChunk)
            currentChunk = []
        currentChunk.append(names[i])

    # Enforce the final chunk to be the same size as the others
    while len(currentChunk) < n:
        currentChunk.append("")

    chunks.append(currentChunk)

    return chunks

def createWorkbook(weeklyOrders):
    """
    Create the weeklyOrders.csv file to be passed off for further processing.
    """
    rows = []
    unique_meals = set()

    for customer in weeklyOrders:
        for i in range(len(weeklyOrders[customer])):
            unique_meals.add(weeklyOrders[customer][i].split("*")[0])

    for customer in weeklyOrders:
        customer_name = customer.split('@')[0]
        address = customer.split('@')[1]
        
        meals = []
        quantities = []

        for i in range(len(weeklyOrders[customer])):
            meals.append(weeklyOrders[customer][i].split("*")[0])
            try:
                quantities.append(int(weeklyOrders[customer][i].split("*")[1].strip(" ")))
            except ValueError:
                pass

        totalMeals = sum(quantities)
        row = [0, customer_name, address, totalMeals]

        for meal in unique_meals:
            if meal in meals:
                for i in range(len(meals)):
                    if meals[i] == meal:
                        try:
                            row.append(quantities[i])
                        except:
                            pass
                        i = len(meals)
            else:
                row.append(0)

        rows.append(row)

    header_row = ["Area Code", "Customer Name", "Address", "Total"]
    for meal in unique_meals:
        header_row.append(meal)

    with open('WeeklyOrders.csv', 'w', newline='') as csvFile:
        writer = csv.writer(csvFile)
        writer.writerow(header_row)
        for row in rows:
            writer.writerow(row)

    csvFile.close()

def main():
    """
    Program entry point.
    """
    with open("user_data/ExtraOrders.txt") as f:
        extraOrders = f.readlines()

    extraOrders = [x.strip() for x in extraOrders]
    
    # Use the default location of user_data/orders.csv unless a different location is set via. commandline.
    # E.g. python parse_order_exports.py test/testOrders.csv
    orderDataFile = "user_data/orders.csv"
    if (len(sys.argv) > 1):
        orderDataFile = str(sys.argv[1])

    customerIDs, weeklyOrders = extractCustomerData(orderDataFile)
    createWorkbook(weeklyOrders)
    customerNames, customerAddresses, customerTotalMeals = generateLabels(customerIDs)
    #createStickers(customerNames, customerAddresses, customerTotalMeals)

if __name__ == '__main__': main()
